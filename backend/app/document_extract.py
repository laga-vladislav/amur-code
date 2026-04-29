"""Extract plain text from PDF/DOCX uploads so they can feed AI generation."""

from __future__ import annotations

import io
from typing import Iterable


MAX_CHARS = 20000


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) <= MAX_CHARS:
        return text
    return text[:MAX_CHARS].rstrip() + "\n…"


def _join_lines(lines: Iterable[str]) -> str:
    cleaned = []
    for line in lines:
        line = (line or "").strip()
        if line:
            cleaned.append(line)
    return "\n".join(cleaned)


def extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            chunks.append(page_text.strip())
    return _truncate("\n\n".join(chunks))


def extract_docx_text(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    return _truncate("\n".join(parts))


def extract_text(filename: str, content_type: str | None, data: bytes) -> str:
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if name.endswith(".pdf") or "pdf" in ct:
        return extract_pdf_text(data)
    if name.endswith(".docx") or "officedocument.wordprocessingml" in ct:
        return extract_docx_text(data)
    if name.endswith(".txt") or ct.startswith("text/"):
        return _truncate(data.decode("utf-8", errors="replace"))
    raise ValueError("Поддерживаются файлы PDF, DOCX и TXT")
